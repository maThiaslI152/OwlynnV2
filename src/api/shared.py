import logging
import re
from langchain_core.messages import AIMessage, ToolMessage

logger = logging.getLogger("src.api")

connected_websockets = set()
_session_usage = {
    "prompt_tokens": 0,
    "completion_tokens": 0,
    "prompt_cache_hit_tokens": 0,
    "prompt_cache_miss_tokens": 0,
    "total_tokens": 0,
}


async def emit_cloud_usage_events(send_ws, *, turn_usage=None, model_used=None):
    """Emit cloud_usage and any newly crossed budget warnings over WebSocket."""
    from src.agent.cloud_cost_tracker import (
        build_cloud_usage_payload,
        get_cost_tracker,
    )
    from src.memory.user_profile import get_profile
    from src.config.config_loader import config

    payload = build_cloud_usage_payload(turn_usage=turn_usage, model_used=model_used)
    await send_ws({"type": "cloud_usage", **payload})

    profile = get_profile()
    daily_limit = int(
        profile.get("cloud_daily_token_limit")
        or config.get("cloud.budget.daily_token_limit", 500_000)
    )
    thresholds = profile.get("cloud_budget_warning_thresholds") or config.get(
        "cloud.budget.warning_thresholds", [0.5, 0.8, 0.95]
    )
    tracker = get_cost_tracker()
    for warning in tracker.consume_budget_warnings(daily_limit, thresholds):
        await send_ws({"type": "cloud_budget_warning", **warning})


_TOOL_DESTRUCTIVE_RE = re.compile(
    r"(?:\brm\s+-rf\b|\bdrop\b|\bdelete\b|\btruncate\b)", re.IGNORECASE
)
_TOOL_NETWORK_RE = re.compile(
    r"(?:\bcurl\b|\bwget\b|\bhttp[s]?://\b|\bscp\b|\bssh\b)", re.IGNORECASE
)
_TOOL_PRIV_RE = re.compile(r"(?:\bsudo\b|\bchmod\b|\bchown\b)", re.IGNORECASE)


def _stringify_lc_message_content(content) -> str:
    """Flatten LangChain message content (str or list of blocks) for JSON/UI."""
    if content is None:
        return ""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts: list[str] = []
        for block in content:
            if isinstance(block, str):
                parts.append(block)
            elif isinstance(block, dict):
                text = block.get("text")
                if text is not None:
                    parts.append(str(text))
                else:
                    nested = block.get("content")
                    if nested is not None:
                        parts.append(_stringify_lc_message_content(nested))
            else:
                parts.append(str(block))
        return "".join(parts)
    return str(content)


def serialize_message(msg):
    """
    Converts Langchain BaseMessage objects into raw UI-friendly dictionaries
    so they can be safely streamed over WebSockets to a React client.
    """
    if isinstance(msg, AIMessage):
        content_ui = _stringify_lc_message_content(msg.content)
        # Strip DSML/Qwen XML tags from assistant messages
        try:
            from src.agent.nodes.complex_utils.formatter import (
                _strip_dsml_blocks,
                _strip_thinking_tags,
            )

            content_ui = _strip_dsml_blocks(_strip_thinking_tags(content_ui or ""))
        except Exception:
            pass
    else:
        content_ui = msg.content

    serialized = {"type": getattr(msg, "type", "unknown"), "content": content_ui}

    if isinstance(msg, AIMessage) and getattr(msg, "tool_calls", None):
        serialized["tool_calls"] = msg.tool_calls

    if isinstance(msg, ToolMessage):
        serialized["tool_name"] = getattr(msg, "name", "unknown")
        serialized["tool_call_id"] = getattr(msg, "tool_call_id", "")
        # Truncate content for UI readability/performance if too large
        if isinstance(msg.content, str) and len(msg.content) > 500:
            serialized["content"] = (
                msg.content[:500] + "\n\n... [Content Truncated for UI] ..."
            )

    return serialized


def extract_pdf_text(raw_bytes: bytes) -> str:
    """Extract text from a PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=raw_bytes, filetype="pdf")
        try:
            pages_text = []
            for page in doc:
                pages_text.append(page.get_text())
        finally:
            doc.close()
        return "\n\n".join(pages_text)
    except Exception as e:
        logger.error("PyMuPDF text extraction failed: %s", e)
        return ""


def extract_docx_text(raw_bytes: bytes) -> str:
    """Extract paragraphs and table text from DOCX bytes using python-docx."""
    import io

    try:
        from docx import Document

        doc = Document(io.BytesIO(raw_bytes))
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            text += "\n--- Table ---\n"
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if row_idx == 0:
                    text += " | ".join(cells) + "\n" + "---" * len(cells) + "\n"
                else:
                    text += " | ".join(cells) + "\n"
        return text
    except Exception as e:
        logger.error("python-docx text extraction failed: %s", e)
        return ""


def extract_doc_text(raw_bytes: bytes) -> str:
    """Extract printable text strings from binary DOC bytes (ASCII & UTF-16-LE)."""
    import re

    # Match ASCII printable characters of length 4 or more
    ascii_pat = re.compile(rb"[\x20-\x7E\x0A\x0D]{4,}")
    # Match UTF-16-LE printable characters (character followed by null byte)
    utf16_pat = re.compile(rb"(?:[\x20-\x7E\x0A\x0D]\x00){4,}")

    strings = []

    # 1. Extract ASCII strings
    for match in ascii_pat.finditer(raw_bytes):
        try:
            s = match.group(0).decode("ascii", errors="ignore").strip()
            if s:
                strings.append(s)
        except Exception:
            pass

    # 2. Extract UTF-16-LE strings
    for match in utf16_pat.finditer(raw_bytes):
        try:
            s = match.group(0).decode("utf-16-le", errors="ignore").strip()
            if s:
                strings.append(s)
        except Exception:
            pass

    # Combine, clean up extra whitespaces/newlines, and join
    combined = "\n".join(strings)
    # Clean up multiple whitespaces
    combined = re.sub(r"[ \t]+", " ", combined)
    # Clean up multiple newlines
    combined = re.sub(r"\n{3,}", "\n\n", combined)
    return combined.strip()


async def build_message_content(text: str, files: list):
    """
    Builds the message content block for LangChain, supporting:
    - Images: forwarded as image_url for multimodal vision models
    - Text PDFs: text extracted via PyMuPDF and injected
    - Word Docs: text extracted via python-docx/binary scanner and injected
    - Scanned PDFs: each page rendered as image and forwarded to the vision model
    - Code/text files: decoded and injected as a fenced code block
    """
    import asyncio

    from src.api.attachment_intake import (
        is_vision_mime,
        lm_studio_safe_image_payload,
        normalize_file_attachment,
    )
    from src.config.config_loader import config

    MAX_INLINE_PDF_CHARS = int(config.get("tool_output.max_inline_pdf_chars", 16000))

    content_parts = []
    text_injections = []
    has_multimodal = False

    for f in files:
        if f.get("type") == "workspace_ref":
            continue
        normalized = normalize_file_attachment(f)
        if not normalized:
            continue

        mime = normalized["type"]
        data_b64 = normalized["data"]
        name = normalized["name"]
        raw_bytes = normalized["raw_bytes"]

        if is_vision_mime(mime):
            has_multimodal = True
            safe_mime, safe_b64 = lm_studio_safe_image_payload(mime, raw_bytes)
            content_parts.append(
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{safe_mime};base64,{safe_b64}"},
                }
            )

        elif mime == "application/pdf" or name.lower().endswith(".pdf"):
            logger.info("Uploaded '%s'. Extracting text for chat context.", name)
            pdf_text = await asyncio.to_thread(extract_pdf_text, raw_bytes)
            pdf_text = (pdf_text or "").strip()
            if len(pdf_text) >= 200:
                excerpt = pdf_text[:MAX_INLINE_PDF_CHARS]
                if len(pdf_text) > MAX_INLINE_PDF_CHARS:
                    excerpt += (
                        "\n\n[PDF truncated in this prompt for size; full file is on disk — "
                        "call read_workspace_file as a real tool if you need the rest.]"
                    )
                text_injections.append(
                    f"[Workspace file `{name}` — text extracted from PDF below. "
                    f"Use this to answer when it is enough; if not, call read_workspace_file with that path "
                    f"(function/tool call, not instructions to the user).]\n\n---\n{excerpt}\n---"
                )
            else:
                text_injections.append(
                    f"[Workspace file `{name}` — little or no extractable text in upload preview. "
                    f"You must invoke read_workspace_file for `{name}` as a tool/function call before answering.]"
                )

        elif name.lower().endswith(".docx"):
            logger.info("Uploaded '%s'. Extracting text for chat context.", name)
            docx_text = await asyncio.to_thread(extract_docx_text, raw_bytes)
            docx_text = (docx_text or "").strip()
            if len(docx_text) >= 50:
                excerpt = docx_text[:MAX_INLINE_PDF_CHARS]
                if len(docx_text) > MAX_INLINE_PDF_CHARS:
                    excerpt += (
                        "\n\n[DOCX truncated in this prompt for size; full file is on disk — "
                        "call read_workspace_file as a real tool if you need the rest.]"
                    )
                text_injections.append(
                    f"[Workspace file `{name}` — text extracted from DOCX below. "
                    f"Use this to answer when it is enough; if not, call read_workspace_file with that path.]\n\n---\n{excerpt}\n---"
                )
            else:
                text_injections.append(
                    f"[Workspace file `{name}` — little or no extractable text in upload preview. "
                    f"You must invoke read_workspace_file for `{name}` as a tool/function call before answering.]"
                )

        elif name.lower().endswith(".doc"):
            logger.info("Uploaded '%s'. Extracting text for chat context.", name)
            doc_text = await asyncio.to_thread(extract_doc_text, raw_bytes)
            doc_text = (doc_text or "").strip()
            if len(doc_text) >= 50:
                excerpt = doc_text[:MAX_INLINE_PDF_CHARS]
                if len(doc_text) > MAX_INLINE_PDF_CHARS:
                    excerpt += (
                        "\n\n[DOC truncated in this prompt for size; full file is on disk — "
                        "call read_workspace_file as a real tool if you need the rest.]"
                    )
                text_injections.append(
                    f"[Workspace file `{name}` — text extracted from DOC below. "
                    f"Use this to answer when it is enough; if not, call read_workspace_file with that path.]\n\n---\n{excerpt}\n---"
                )
            else:
                text_injections.append(
                    f"[Workspace file `{name}` — little or no extractable text in upload preview. "
                    f"You must invoke read_workspace_file for `{name}` as a tool/function call before answering.]"
                )

        else:
            # Text / code file
            try:
                decoded_text = raw_bytes.decode("utf-8")
                logger.info("Uploaded '%s'. Inlining text directly into context.", name)
                text_injections.append(
                    f"[Attached File: `{name}`]\n```\n{decoded_text}\n```\n"
                )
            except UnicodeDecodeError:
                logger.info(
                    "Uploaded '%s' is binary. Adding workspace reference.", name
                )
                text_injections.append(
                    f"[Workspace file `{name}` saved. Invoke read_workspace_file as a tool with that path if you need contents — "
                    f"do not answer with only a suggestion to use the tool.]"
                )

    # Build final content
    if has_multimodal:
        # Multimodal message — prepend any text file injections
        for inj in text_injections:
            content_parts.insert(0, {"type": "text", "text": inj})
        if text:
            content_parts.append({"type": "text", "text": text})
        return content_parts if content_parts else None
    else:
        # Plain text message
        parts = text_injections[:]
        if text:
            parts.append(text)
        return "\n\n".join(parts) if parts else None
