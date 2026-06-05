"""Generate controlled sample files for the RAG file intake audit.

Each file contains known, verifiable content that we can test against
after the LLM reads them via the Owlynn pipeline.
"""

import os

SAMPLES_DIR = os.path.dirname(os.path.abspath(__file__))


def create_pdf():
    """2-page mixed document with a unique revenue stat."""
    import fitz  # PyMuPDF

    doc = fitz.Document()
    page1 = doc.new_page()

    # Insert page 1 content immediately (before creating page 2)
    html1 = (
        '<p style="font-size:18px;"><b>Owlynn RAG Audit - Test Document PDF</b></p>'
        "<br>"
        '<p style="font-size:14px;"><b>Project Omega - Q1 2026 Financial Summary</b></p>'
        "<hr>"
        '<p style="font-size:13px;"><b>Project Omega revenue: 4,200,000 THB (4.2M) in Q1 2026</b></p>'
        '<p style="font-size:12px;">Gross margin: 62 percent  |  Operating costs: 1,600,000 THB</p>'
        '<p style="font-size:12px;">Team size: 12 engineers  |  Office: Bangkok, Thailand</p>'
    )
    page1.insert_htmlbox(fitz.Rect(50, 50, 545, 500), html1)

    # Now create page 2 and insert its content
    page2 = doc.new_page()
    html2 = (
        '<p style="font-size:16px;"><b>Page 2 - Additional Details</b></p>'
        "<br>"
        '<p style="font-size:12px;">Project Omega launched in January 2026 with a focus</p>'
        '<p style="font-size:12px;">on AI-powered document intelligence. By end of Q1,</p>'
        '<p style="font-size:12px;">the platform processed 15,000+ documents across 3 formats.</p>'
        "<br>"
        '<p style="font-size:13px;"><b>Key Technologies:</b></p>'
        "<p>1. PyMuPDF for PDF text extraction</p>"
        "<p>2. python-docx for Word documents</p>"
        "<p>3. pandas + openpyxl for Excel spreadsheets</p>"
    )
    page2.insert_htmlbox(fitz.Rect(50, 50, 545, 500), html2)

    path = os.path.join(SAMPLES_DIR, "audit_doc.pdf")
    doc.save(path)
    doc.close()
    print(f"Created {path}")
    return path


def create_docx():
    """Word document with a named list of 5 items + a spec table."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Audit Specification Document", level=1)
    doc.add_paragraph(
        "This document contains the specification for the Owlynn RAG audit test."
    )

    doc.add_heading("Core Components (5 items)", level=2)
    items = [
        "1. File Watcher — Auto-detects new files in workspace/",
        "2. Text Extractor — Converts PDF/DOCX/XLSX to plain text",
        "3. Processed Cache — Stores extracted text in .processed/",
        "4. Qdrant Indexer — Vectorizes and indexes document chunks",
        "5. RAG Search Tool — Semantic retrieval via search_workspace_docs",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Technical Specifications", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Library"
    hdr[2].text = "Version"

    specs = [
        ("PDF Parser", "PyMuPDF (fitz)", ">=1.23"),
        ("DOCX Parser", "python-docx", ">=1.0"),
        ("XLSX Parser", "pandas + openpyxl", ">=3.0"),
        ("Vector Store", "Qdrant via Mem0", ">=1.0"),
        ("LLM Engine", "Gemma 4 E4B (LM Studio)", "Q4_K_M"),
    ]
    for comp, lib, ver in specs:
        row = table.add_row().cells
        row[0].text = comp
        row[1].text = lib
        row[2].text = ver

    path = os.path.join(SAMPLES_DIR, "audit_doc.docx")
    doc.save(path)
    print(f"Created {path}")
    return path


def create_xlsx():
    """Excel spreadsheet with Product, Quantity, Price for 6 rows."""
    import openpyxl
    from openpyxl.styles import Font

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Inventory Q1 2026"

    # Header
    headers = ["Product", "Quantity", "Price (THB)"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True)

    # Data rows — known values for verification
    data = [
        ("Wireless Mouse", 145, 890),
        ("Mechanical Keyboard", 78, 3200),
        ("USB-C Hub", 210, 1250),
        ("27-inch Monitor", 32, 8900),
        ("Webcam 4K", 56, 4500),
        ("Noise-Canceling Headphones", 89, 5800),
    ]
    for row_idx, (product, qty, price) in enumerate(data, 2):
        ws.cell(row=row_idx, column=1, value=product)
        ws.cell(row=row_idx, column=2, value=qty)
        ws.cell(row=row_idx, column=3, value=price)

    # Auto-adjust column widths
    for col in ws.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = max_len + 4

    path = os.path.join(SAMPLES_DIR, "audit_data.xlsx")
    wb.save(path)
    print(f"Created {path}")
    return path


if __name__ == "__main__":
    create_pdf()
    create_docx()
    create_xlsx()
    print("\nAll sample files created successfully.")
