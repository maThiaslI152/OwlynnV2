"""Tests for DOCX and DOC text extraction."""

import os
import sys
import tempfile
from unittest.mock import MagicMock

sys.modules["mem0"] = MagicMock()

import pytest
from src.api.file_processor import FileWatcherHandler
from src.api.shared import extract_docx_text, extract_doc_text


def test_docx_extraction():
    # 1. Create a dummy docx file using python-docx
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello world, this is a test document.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Val 1"
    table.cell(1, 1).text = "Val 2"

    with tempfile.TemporaryDirectory() as temp_dir:
        filepath = os.path.join(temp_dir, "test.docx")
        doc.save(filepath)

        # 2. Test shared extractor
        with open(filepath, "rb") as f:
            raw_bytes = f.read()
        extracted = extract_docx_text(raw_bytes)
        assert "Hello world" in extracted
        assert "Header 1" in extracted
        assert "Val 2" in extracted

        # 3. Test FileWatcherHandler
        handler = FileWatcherHandler(temp_dir)
        handler.process_file(filepath)

        # Verify output
        output_path = os.path.join(temp_dir, ".processed", "test.docx.txt")
        assert os.path.exists(output_path), "DOCX output not created in processed dir"
        with open(output_path, "r", encoding="utf-8") as f:
            content = f.read()
        assert "Hello world" in content
        assert "Header 1" in content
        assert "Val 2" in content


def test_doc_binary_extraction():
    # Create a simulated binary doc content with ASCII and UTF-16-LE strings
    # Must use strings of length 4 or more to match patterns
    content = (
        b"Some random garbage "
        + b"Important ASCIIText here"
        + b" more garbage "
        + "Important UTF16Text".encode("utf-16-le")
        + b" even more garbage"
    )

    extracted = extract_doc_text(content)
    assert "ASCIIText" in extracted
    assert "UTF16Text" in extracted

    # Write to a simulated .doc file
    with tempfile.TemporaryDirectory() as temp_dir:
        filepath = os.path.join(temp_dir, "test_legacy.doc")
        with open(filepath, "wb") as f:
            f.write(content)

        # Test FileWatcherHandler
        handler = FileWatcherHandler(temp_dir)
        handler.process_file(filepath)

        # Verify output
        output_path = os.path.join(temp_dir, ".processed", "test_legacy.doc.txt")
        assert os.path.exists(output_path), "DOC output not created in processed dir"
        with open(output_path, "r", encoding="utf-8") as f:
            processed_content = f.read()
        assert "ASCIIText" in processed_content
        assert "UTF16Text" in processed_content
